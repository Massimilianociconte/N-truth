"""End-to-end ingest, block segmentation and coreference regressions.

These tests use real container formats generated in a temporary directory.  They
exercise Project.add -> parser registry -> Document IR -> analysis, rather than
calling parser helpers with synthetic RawDocument objects.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from ntruth.graph.index import GraphIndex
from ntruth.ingest.project import Project
from ntruth.ingest.safety import SafetyError
from ntruth.parsers.registry import build_document_ir
from ntruth.pipeline import analyze_project
from ntruth.schemas.document import ParserStatus, SectionRole
from ntruth.schemas.graph import NodeType, RelationType


def _ingest(source: Path, workspace: Path, *, domain: str = "quantitative_microscopy") -> Project:
    project = Project.create(workspace, name=source.stem, domain=domain, language="en")
    result = project.add(source)
    assert result.accepted, result.summary()
    return project


def _write_text_pdf(path: Path, text: str) -> None:
    """Write a small real PDF whose text is extractable by pypdf."""
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
    )
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = DecodedStreamObject()
    stream.set_data(f"BT /F1 11 Tf 40 740 Td ({escaped}) Tj ET".encode("latin-1"))
    page[NameObject("/Contents")] = writer._add_object(stream)
    with path.open("wb") as output:
        writer.write(output)


def test_two_explicit_experiments_are_isolated_into_distinct_blocks(tmp_path: Path) -> None:
    source = tmp_path / "two-experiments.md"
    source.write_text(
        """# Experiment 1

## Materials and Methods

Cells were prepared from three independent cultures. Cells were treated with drug or
vehicle at the level of the culture. n = 120 cells.

# Experiment 2

## Materials and Methods

Cells were prepared from five independent cultures. Cells were treated with drug or
vehicle at the level of the culture. n = 200 cells.
""",
        encoding="utf-8",
    )

    result = analyze_project(_ingest(source, tmp_path / "project"))

    assert [block.title for block in result.report.blocks] == ["Experiment 1", "Experiment 2"]
    assert [
        {statement.value for statement in block.n_statements} for block in result.report.blocks
    ] == [
        {120},
        {200},
    ]
    assert all(not block.contradictions for block in result.report.blocks)


def test_plain_text_experiment_headings_are_structural(tmp_path: Path) -> None:
    source = tmp_path / "two-experiments.txt"
    source.write_text(
        """Experiment 1
Methods
Three cultures were prepared; n = 3 cultures.

Experiment 2
Methods
Five cultures were prepared; n = 5 cultures.
""",
        encoding="utf-8",
    )

    result = analyze_project(_ingest(source, tmp_path / "project"))

    assert [block.title for block in result.report.blocks] == ["Experiment 1", "Experiment 2"]
    assert [{n.value for n in block.n_statements} for block in result.report.blocks] == [
        {3},
        {5},
    ]


def test_prose_study_headings_do_not_manufacture_blocks(tmp_path: Path) -> None:
    source = tmp_path / "study.txt"
    source.write_text(
        """Study design
Three cultures were prepared.

Study population
The cultures were sampled once.
""",
        encoding="utf-8",
    )

    result = analyze_project(_ingest(source, tmp_path / "project"))

    assert len(result.report.blocks) == 1


def test_deterministic_coreference_is_explicit_and_block_local(tmp_path: Path) -> None:
    source = tmp_path / "coreference.md"
    source.write_text(
        """# Experiment 1

## Materials and Methods

Three independent cultures were prepared. These cultures were treated with drug or vehicle.

# Experiment 2

## Materials and Methods

Five independent cultures were prepared. Those cultures were treated with drug or vehicle.
""",
        encoding="utf-8",
    )

    result = analyze_project(_ingest(source, tmp_path / "project"))

    assert len(result.report.blocks) == 2
    for block in result.report.blocks:
        assert len(block.coreference_links) == 1
        link = block.coreference_links[0]
        mention_ids = {mention.id for mention in block.mentions}
        assert link.anaphor_mention_id in mention_ids
        assert link.antecedent_mention_id in mention_ids
        anaphor = next(m for m in block.mentions if m.id == link.anaphor_mention_id)
        assert anaphor.is_anaphor


def test_tampered_manifest_cannot_escape_project_sources(tmp_path: Path) -> None:
    source = tmp_path / "legit.txt"
    source.write_text("Methods\n\nlegitimate content", encoding="utf-8")
    outside = tmp_path / "outside.txt"
    outside.write_text("Methods\n\nOUTSIDE_WORKSPACE", encoding="utf-8")
    project = _ingest(source, tmp_path / "project")
    project_file = project.manifest.files[0].model_copy(
        update={
            "relative_path": "../outside.txt",
            "filename": "outside.txt",
            "size_bytes": outside.stat().st_size,
        }
    )
    project.manifest = project.manifest.model_copy(update={"files": (project_file,)})
    project.save()

    with pytest.raises(SafetyError, match=r"fuori|sources"):
        Project.open(project.root)


def test_xlsx_formulas_are_preserved_as_inert_text_and_flagged(tmp_path: Path) -> None:
    source = tmp_path / "samples.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    assert sheet is not None
    sheet.title = "Samples"
    sheet.append(["sample_id", "donor", "computed"])
    sheet.append(["S1", "D1", "=1+1"])
    workbook.save(source)

    ir = build_document_ir(_ingest(source, tmp_path / "project"))

    assert ir.tables[0].cell(0, "computed") == "'=1+1"
    assert "formula" in " ".join(ir.files[0].warnings)


def test_xlsx_formula_headers_are_inert_and_flagged(tmp_path: Path) -> None:
    source = tmp_path / "formula-header.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    assert sheet is not None
    sheet.append(['=HYPERLINK("https://invalid.example","click")', "donor"])
    sheet.append(["payload", "D1"])
    workbook.save(source)

    ir = build_document_ir(_ingest(source, tmp_path / "project"))

    assert ir.tables[0].columns[0].startswith("'=")
    assert "formula" in " ".join(ir.files[0].warnings)


def test_real_docx_reaches_document_ir_and_analysis(tmp_path: Path) -> None:
    source = tmp_path / "methods.docx"
    document = Document()
    document.add_heading("Materials and Methods", level=1)
    document.add_paragraph("Cells were quantified in three wells; n = 3 wells.")
    document.save(str(source))

    project = _ingest(source, tmp_path / "project")
    result = analyze_project(project)

    assert result.document.files[0].parser == "docx"
    assert result.document.files[0].status is ParserStatus.OK
    assert any(statement.value == 3 for statement in result.block.n_statements)


def test_real_xlsx_multiple_sheets_reach_document_ir_and_analysis(tmp_path: Path) -> None:
    source = tmp_path / "samples.xlsx"
    workbook = Workbook()
    first = workbook.active
    assert first is not None
    first.title = "Experiment 1"
    first.append(["experiment_id", "donor", "well", "treatment"])
    first.append(["1", "D1", "W1", "drug"])
    first.append(["1", "D2", "W2", "vehicle"])
    second = workbook.create_sheet("Experiment 2")
    second.append(["experiment_id", "donor", "well", "treatment"])
    second.append(["2", "D3", "W3", "drug"])
    second.append(["2", "D4", "W4", "vehicle"])
    workbook.save(source)

    result = analyze_project(_ingest(source, tmp_path / "project"))

    assert result.document.files[0].parser == "xlsx"
    assert {table.sheet for table in result.document.tables} == {"Experiment 1", "Experiment 2"}
    assert all(table.rows for table in result.document.tables)


def test_sample_sheet_builds_instance_graph_and_keeps_aggregate_adapter(
    tmp_path: Path,
) -> None:
    source = tmp_path / "samples.csv"
    source.write_text(
        "sample_id,preparation,well,treatment\n"
        "S1,P1,W1,drug\n"
        "S2,P1,W2,drug\n"
        "S3,P1,W3,vehicle\n"
        "S4,P1,W4,vehicle\n",
        encoding="utf-8",
    )

    result = analyze_project(_ingest(source, tmp_path / "project"))
    index = GraphIndex(result.block.hierarchy)

    aggregate_well = index.node(NodeType.WELL)
    assert aggregate_well is not None
    assert aggregate_well.attributes["aggregate"] is True
    assert index.count(NodeType.WELL) == 4
    assert len(index.instances(NodeType.WELL)) == 4
    assert len(index.instances(NodeType.CELL_CULTURE)) == 1
    assert index.scoped_instance_count(NodeType.WELL, factor_name="treatment", group="drug") == 2

    instance_ids = {node.id for node in index.instances(NodeType.WELL)}
    assert all("W1" not in node_id for node_id in instance_ids)
    nested = [
        relation
        for relation in result.block.hierarchy.relations
        if relation.type is RelationType.NESTED_IN and relation.attributes.get("instance_relation")
    ]
    assert len(nested) == 4
    assert all(relation.source in instance_ids for relation in nested)


def test_single_sample_sheet_is_split_by_explicit_experiment_id(tmp_path: Path) -> None:
    source = tmp_path / "samples.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    assert sheet is not None
    sheet.title = "Samples"
    sheet.append(["experiment_id", "donor", "treatment"])
    sheet.append(["1", "D1", "drug"])
    sheet.append(["1", "D2", "vehicle"])
    sheet.append(["2", "D3", "drug"])
    sheet.append(["2", "D4", "vehicle"])
    workbook.save(source)

    result = analyze_project(_ingest(source, tmp_path / "project"))

    assert [block.title for block in result.report.blocks] == ["Experiment 1", "Experiment 2"]
    rows_by_block = [analysis.document.tables[0].rows for analysis in result.block_analyses]
    assert [{row["experiment_id"] for row in rows} for rows in rows_by_block] == [
        {"1"},
        {"2"},
    ]


def test_real_text_pdf_reaches_document_ir_with_quality_status(tmp_path: Path) -> None:
    source = tmp_path / "methods.pdf"
    sentence = "Methods. Cells were quantified in three wells; n = 3 wells. "
    _write_text_pdf(source, sentence * 5)

    project = _ingest(source, tmp_path / "project")
    result = analyze_project(project)

    assert result.document.files[0].parser == "pdf"
    assert result.document.files[0].status is ParserStatus.OK
    assert any(statement.value == 3 for statement in result.block.n_statements)


def test_real_jats_preserves_legend_table_and_n_statement(tmp_path: Path) -> None:
    source = tmp_path / "article.nxml"
    source.write_text(
        """<article>
  <front><article-meta><title-group><article-title>Study</article-title></title-group></article-meta></front>
  <body><sec><title>Methods</title><p>Cells were quantified.</p></sec></body>
  <fig id="f1"><label>Figure 1</label><caption><p>Quantification; n = 12 cells.</p></caption></fig>
  <table-wrap id="t1"><label>Table 1</label><caption><p>Samples</p></caption>
    <table><tr><th>donor</th><th>well</th></tr><tr><td>D1</td><td>W1</td></tr></table>
  </table-wrap>
</article>""",
        encoding="utf-8",
    )

    result = analyze_project(_ingest(source, tmp_path / "project"))

    assert result.document.files[0].parser == "jats"
    assert result.document.tables and result.document.tables[0].caption
    assert SectionRole.FIGURE_LEGEND in {section.role for section in result.document.sections}
    assert any(statement.value == 12 for statement in result.block.n_statements)


def test_low_density_pdf_challenge_is_degraded_and_forces_abstention(tmp_path: Path) -> None:
    source = tmp_path / "ocr-like.pdf"
    _write_text_pdf(source, "n = 3 wells")

    result = analyze_project(_ingest(source, tmp_path / "project"))

    assert result.document.files[0].status is ParserStatus.DEGRADED
    assert result.abstention.abstained
    assert not [alert for alert in result.block.alerts if alert.severity.value == "critical"]


def test_out_of_domain_challenge_is_visible_in_report_limits(tmp_path: Path) -> None:
    source = tmp_path / "methods.txt"
    source.write_text("Methods\n\nThree observations were reported.", encoding="utf-8")

    result = analyze_project(_ingest(source, tmp_path / "project", domain="unvalidated_assay"))

    assert any("fuori dal perimetro" in limit.lower() for limit in result.report.limits)


def test_sample_sheet_missing_ids_challenge_is_visible(tmp_path: Path) -> None:
    source = tmp_path / "samples.csv"
    source.write_text(
        "sample_id,donor,well,treatment\nS1,D1,W1,drug\nS2,,W2,vehicle\n",
        encoding="utf-8",
    )

    result = analyze_project(_ingest(source, tmp_path / "project"))

    assert any("valori mancanti" in limit.lower() for limit in result.report.limits)
