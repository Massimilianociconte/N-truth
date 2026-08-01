"""Parser DOCX (PRD FR-003). Round-trip verificato dalle fixture."""

from __future__ import annotations

from pathlib import Path

from ntruth.parsers.base import ParseFailure, RawBlock, RawDocument, RawTable
from ntruth.parsers.sections import looks_like_heading
from ntruth.schemas.document import ParserStatus


class DocxParser:
    name = "docx"

    def supports(self, path: Path, media_type: str) -> bool:
        return path.suffix.lower() == ".docx"

    def parse(self, path: Path) -> RawDocument:
        try:
            import docx  # python-docx
        except ImportError as exc:  # pragma: no cover - dipendenza dichiarata
            raise ParseFailure(path, "python-docx non installato") from exc

        try:
            document = docx.Document(str(path))
        except Exception as exc:
            raise ParseFailure(path, f"DOCX illeggibile ({type(exc).__name__})") from exc

        doc = RawDocument(parser=self.name)
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = (paragraph.style.name or "").lower() if paragraph.style else ""
            if style.startswith("heading") or style in {"title", "subtitle"}:
                level = _heading_level_from_style(style)
                doc.blocks.append(RawBlock(kind="heading", text=text, level=level))
            elif style.startswith("caption"):
                doc.blocks.append(RawBlock(kind="caption", text=text))
            elif looks_like_heading(text) and len(text.split()) <= 12:
                doc.blocks.append(RawBlock(kind="heading", text=text, level=2))
            else:
                doc.blocks.append(RawBlock(kind="paragraph", text=text))

        for index, table in enumerate(document.tables):
            parsed = _convert_table(table, f"table-{index + 1}")
            if parsed is not None:
                doc.tables.append(parsed)

        if doc.is_empty:
            doc.status = ParserStatus.FAILED
            doc.warnings.append("DOCX senza paragrafi ne tabelle leggibili")
        return doc


def _heading_level_from_style(style: str) -> int:
    if style in {"title"}:
        return 1
    digits = "".join(ch for ch in style if ch.isdigit())
    return int(digits) if digits else 2


def _convert_table(table: object, name: str) -> RawTable | None:
    rows = getattr(table, "rows", [])
    if not rows:
        return None
    header_cells = [c.text.strip() for c in rows[0].cells]
    header = [h or f"col_{i + 1}" for i, h in enumerate(header_cells)]
    out = RawTable(name=name, columns=header)
    for row in rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < len(header):
            cells.extend([""] * (len(header) - len(cells)))
        out.rows.append(dict(zip(header, cells[: len(header)], strict=True)))
    return out
